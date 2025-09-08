import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import geopandas as gpd
import math
from io import BytesIO
import re

# Custom CSS for the dashboard
st.markdown("""
<style>
    .main .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
        max-width: none;
    }
    
    h1 {
        color: #2c3e50;
        text-align: center;
        font-weight: 700;
        margin-bottom: 2rem;
    }
    
    h2, h3 {
        color: #34495e;
        border-bottom: 2px solid #3498db;
        padding-bottom: 0.5rem;
    }
    
    .stButton > button {
        background: linear-gradient(45deg, #3498db, #2980b9);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 2rem;
        font-weight: 600;
    }
</style>
""", unsafe_allow_html=True)

def create_chiefdom_mapping():
    """Create mapping between GPS data chiefdom names and shapefile FIRST_CHIE names"""
    chiefdom_mapping = {
        # BO District mappings
        "Bo City": "BO TOWN",
        "Badjia": "BADJIA",
        "Bargbo": "BAGBO",
        "Bagbwe": "BAGBWE(BAGBE)",
        "Baoma": "BOAMA",
        "Bongor": "BONGOR",
        "Bumpeh": "BUMPE NGAO",
        "Gbo": "GBO",
        "Jaiama": "JAIAMA",
        "Kakua": "KAKUA",
        "Komboya": "KOMBOYA",
        "Lugbu": "LUGBU",
        "Niawa Lenga": "NIAWA LENGA",
        "Selenga": "SELENGA",
        "Tinkoko": "TIKONKO",
        "Valunia": "VALUNIA",
        "Wonde": "WONDE",
        
        # BOMBALI District mappings
        "Biriwa": "BIRIWA",
        "Bombali Sebora": "BOMBALI SEBORA",
        "Bombali Serry": "BOMBALI SIARI",
        "Gbanti (Bombali)": "GBANTI",
        "Gbanti": "GBANTI",
        "Gbendembu": "GBENDEMBU",
        "Kamaranka": "KAMARANKA",
        "Magbaimba Ndohahun": "MAGBAIMBA NDORWAHUN",
        "Makarie": "MAKARI",
        "Mara": "MARA",
        "Ngowahun": "N'GOWAHUN",
        "Paki Masabong": "PAKI MASABONG",
        "Safroko Limba": "SAFROKO LIMBA",
        "Makeni City": "MAKENI CITY",
    }
    return chiefdom_mapping

def map_chiefdom_name(chiefdom_name, mapping):
    """Map chiefdom name from GPS data to shapefile name"""
    if pd.isna(chiefdom_name):
        return None
    
    chiefdom_name = str(chiefdom_name).strip()
    
    # Direct match
    if chiefdom_name in mapping:
        return mapping[chiefdom_name]
    
    # Case-insensitive match
    for key, value in mapping.items():
        if key.upper() == chiefdom_name.upper():
            return value
    
    # Partial match (contains)
    for key, value in mapping.items():
        if key.upper() in chiefdom_name.upper() or chiefdom_name.upper() in key.upper():
            return value
    
    # Return original if no mapping found
    return chiefdom_name

def extract_itn_data_from_excel(df):
    """Extract ITN coverage data from the Excel file"""
    # Create empty lists to store extracted data
    districts, chiefdoms, total_enrollment, total_itns, distributed_itns = [], [], [], [], []
    
    # Get chiefdom mapping
    chiefdom_mapping = create_chiefdom_mapping()
    
    # Process each row in the "Scan QR code" column
    for idx, qr_text in enumerate(df["Scan QR code"]):
        if pd.isna(qr_text):
            districts.append(None)
            chiefdoms.append(None)
            total_enrollment.append(0)
            total_itns.append(0)
            distributed_itns.append(0)
            continue
            
        # Extract values using regex patterns
        district_match = re.search(r"District:\s*([^\n]+)", str(qr_text))
        district = district_match.group(1).strip() if district_match else None
        districts.append(district)
        
        chiefdom_match = re.search(r"Chiefdom:\s*([^\n]+)", str(qr_text))
        original_chiefdom = chiefdom_match.group(1).strip() if chiefdom_match else None
        
        # Map chiefdom name to match shapefile
        mapped_chiefdom = map_chiefdom_name(original_chiefdom, chiefdom_mapping)
        chiefdoms.append(mapped_chiefdom)
        
        # Calculate total enrollment (sum of all class enrollments)
        enrollment_total = 0
        for class_num in range(1, 6):  # Classes 1-5
            enrollment_col = f"How many pupils are enrolled in Class {class_num}?"
            if enrollment_col in df.columns:
                class_enrollment = df[enrollment_col].iloc[idx]
                if pd.notna(class_enrollment):
                    enrollment_total += int(class_enrollment)
        
        total_enrollment.append(enrollment_total)
        
        # Calculate total ITNs distributed (boys + girls + left at school)
        itns_boys_total = 0
        itns_girls_total = 0
        itns_left_total = 0
        
        for class_num in range(1, 6):  # Classes 1-5
            # ITNs distributed to boys
            boys_col = f"How many boys in Class {class_num} received ITNs?"
            if boys_col in df.columns:
                boys_itns = df[boys_col].iloc[idx]
                if pd.notna(boys_itns):
                    itns_boys_total += int(boys_itns)
            
            # ITNs distributed to girls
            girls_col = f"How many girls in Class {class_num} received ITNs?"
            if girls_col in df.columns:
                girls_itns = df[girls_col].iloc[idx]
                if pd.notna(girls_itns):
                    itns_girls_total += int(girls_itns)
        
        # ITNs left at the school (this appears to be a single value per school, not per class)
        left_col = "ITNs left at the school for pupils who were absent."
        if left_col in df.columns:
            left_itns = df[left_col].iloc[idx]
            if pd.notna(left_itns):
                itns_left_total = int(left_itns)
        
        # Total ITNs = boys + girls + left at school
        itns_total = itns_boys_total + itns_girls_total + itns_left_total
        itns_distributed = itns_boys_total + itns_girls_total + itns_left_total
        
        total_itns.append(itns_total)
        distributed_itns.append(itns_distributed)
    
    # Create a new DataFrame with extracted values
    itn_df = pd.DataFrame({
        "District": districts,
        "Chiefdom": chiefdoms,
        "Total_Enrollment": total_enrollment,
        "Total_ITNs": total_itns,
        "Distributed_ITNs": distributed_itns
    })
    
    return itn_df

def get_coverage_color(coverage_percent):
    """Get color based on coverage percentage"""
    if coverage_percent < 20:
        return '#d32f2f'  # Red
    elif coverage_percent < 40:
        return '#f57c00'  # Orange
    elif coverage_percent < 60:
        return '#fbc02d'  # Yellow
    elif coverage_percent < 80:
        return '#388e3c'  # Light Green
    elif coverage_percent < 100:
        return '#1976d2'  # Blue
    else:
        return '#4a148c'  # Purple (100% coverage)

def create_itn_coverage_dashboard(gdf, itn_df, district_name, cols=4):
    """Create ITN coverage dashboard optimized for Word document export"""
    
    # Filter shapefile for the district
    district_gdf = gdf[gdf['FIRST_DNAM'] == district_name].copy()
    
    if len(district_gdf) == 0:
        st.error(f"No chiefdoms found for {district_name} district in shapefile")
        return None
    
    # Get unique chiefdoms from shapefile
    chiefdoms = sorted(district_gdf['FIRST_CHIE'].dropna().unique())
    
    # Calculate rows needed
    rows = math.ceil(len(chiefdoms) / cols)
    
    # Optimize figure size for Word document (16:10 aspect ratio works well)
    fig_width = 16  # Width for Word document
    fig_height = rows * 3.5  # Height per row optimized for Word
    
    # Create subplot figure optimized for Word export
    fig, axes = plt.subplots(rows, cols, figsize=(fig_width, fig_height))
    fig.suptitle(f'{district_name} District - ITN Coverage Analysis', 
                 fontsize=18, fontweight='bold', y=0.98)
    
    # Ensure axes is always 2D array
    if rows == 1:
        axes = axes.reshape(1, -1)
    elif cols == 1:
        axes = axes.reshape(-1, 1)
    
    # Plot each chiefdom
    for idx, chiefdom in enumerate(chiefdoms):
        row = idx // cols
        col = idx % cols
        ax = axes[row, col]
        
        # Filter shapefile for this specific chiefdom
        chiefdom_gdf = district_gdf[district_gdf['FIRST_CHIE'] == chiefdom].copy()
        
        # Filter ITN data for this district and chiefdom
        district_data = itn_df[itn_df["District"].str.upper() == district_name.upper()].copy()
        chiefdom_data = district_data[district_data["Chiefdom"] == chiefdom].copy()
        
        # Calculate totals for this chiefdom
        enrollment_total = int(chiefdom_data["Total_Enrollment"].sum()) if len(chiefdom_data) > 0 else 0
        itns_total = int(chiefdom_data["Total_ITNs"].sum()) if len(chiefdom_data) > 0 else 0
        
        # Calculate coverage percentage
        coverage_percent = (itns_total / enrollment_total * 100) if enrollment_total > 0 else 0
        coverage_percent = min(coverage_percent, 100)  # Cap at 100%
        
        # Get color based on coverage
        coverage_color = get_coverage_color(coverage_percent)
        
        # Plot chiefdom boundary with coverage color
        chiefdom_gdf.plot(ax=ax, color=coverage_color, edgecolor='black', alpha=0.8, linewidth=1.5)
        
        # Create ITN coverage text (n, m) format
        itn_text = f"({itns_total}, {enrollment_total})"
        
        # Set title with ITN coverage information (optimized font size for Word)
        ax.set_title(f'{chiefdom}\n{itn_text}', 
                    fontsize=10, fontweight='bold', pad=8)
        
        # Add coverage percentage in the center of the chiefdom
        if len(chiefdom_gdf) > 0:
            # Get center of chiefdom
            bounds = chiefdom_gdf.total_bounds
            center_x = (bounds[0] + bounds[2]) / 2
            center_y = (bounds[1] + bounds[3]) / 2
            
            # Add coverage percentage text in the center
            ax.text(center_x, center_y, f"{coverage_percent:.0f}%", 
                   fontsize=14, fontweight='bold', color='white', 
                   ha='center', va='center',
                   bbox=dict(boxstyle='round,pad=0.5', facecolor='black', alpha=0.7))
        
        # Remove axis labels and ticks for cleaner look
        ax.set_xticks([])
        ax.set_yticks([])
        ax.set_xlabel('')
        ax.set_ylabel('')
        
        # Remove the box frame
        for spine in ax.spines.values():
            spine.set_visible(False)
        
        # Add very light grid
        ax.grid(True, alpha=0.2, linestyle='--', linewidth=0.5)
        
        # Set equal aspect ratio
        ax.set_aspect('equal')
        
        # Set bounds to chiefdom extent with minimal padding for better fit
        bounds = chiefdom_gdf.total_bounds
        padding = 0.005  # Reduced padding for better fit in Word
        ax.set_xlim(bounds[0] - padding, bounds[2] + padding)
        ax.set_ylim(bounds[1] - padding, bounds[3] + padding)
    
    # Hide empty subplots
    total_plots = rows * cols
    for idx in range(len(chiefdoms), total_plots):
        row = idx // cols
        col = idx % cols
        axes[row, col].set_visible(False)
    
    # Optimize layout for Word document
    plt.tight_layout()
    plt.subplots_adjust(top=0.90, hspace=0.35, wspace=0.25)  # Increased space below title
    
    return fig

def generate_simple_summary(itn_df):
    """Generate simple summary with just totals and coverage"""
    
    summary_data = []
    
    # District totals
    for district in ["BO", "BOMBALI"]:
        district_data = itn_df[itn_df["District"].str.upper() == district.upper()]
        
        total_enrollment = int(district_data["Total_Enrollment"].sum())
        total_itns_distributed = int(district_data["Distributed_ITNs"].sum())
        coverage = (total_itns_distributed / total_enrollment * 100) if total_enrollment > 0 else 0
        
        summary_data.append({
            'Level': 'District',
            'Name': district,
            'Total_Enrollment': total_enrollment,
            'ITNs_Distributed': total_itns_distributed,
            'Coverage': f"{coverage:.1f}%"
        })
    
    # Chiefdom totals for each district
    for district in ["BO", "BOMBALI"]:
        district_data = itn_df[itn_df["District"].str.upper() == district.upper()]
        
        # Group by chiefdom
        chiefdoms = district_data['Chiefdom'].dropna().unique()
        
        for chiefdom in sorted(chiefdoms):
            chiefdom_data = district_data[district_data['Chiefdom'] == chiefdom]
            
            total_enrollment = int(chiefdom_data["Total_Enrollment"].sum())
            total_itns_distributed = int(chiefdom_data["Distributed_ITNs"].sum())
            coverage = (total_itns_distributed / total_enrollment * 100) if total_enrollment > 0 else 0
            
            summary_data.append({
                'Level': f'{district} Chiefdom',
                'Name': chiefdom,
                'Total_Enrollment': total_enrollment,
                'ITNs_Distributed': total_itns_distributed,
                'Coverage': f"{coverage:.1f}%"
            })
    
    return summary_data

# Streamlit App
st.title("🛡️ Section 3: ITN Coverage Analysis")
st.markdown("**ITN distribution effectiveness by chiefdom**")

# Color legend
st.markdown("""
### Coverage Color Legend:
- 🔴 **Red**: < 20%
- 🟠 **Orange**: 20-39%
- 🟡 **Yellow**: 40-59%
- 🟢 **Light Green**: 60-79%
- 🔵 **Blue**: 80-99%
- 🟣 **Purple**: 100%+
""")

# Coverage format explanation
st.markdown("""
### ITN Coverage Format:
- **Title**: `(ITNs Distributed, Total Enrollment)`
- **Center**: Coverage percentage
- **Colors**: Same as above color legend
- **ITN Calculation**: Boys ITNs + Girls ITNs + ITNs Left at School
""")

# File Information
st.info("""
**📁 Embedded Files:** `sbd first_submission_clean.xlsx` | `Chiefdom2021.shp`  
**📊 Layout:** Fixed 4-column grid optimized for Word export
""")

# Load the embedded data files
try:
    # Load Excel file (embedded)
    df_original = pd.read_excel("SBD_Final_data_dissemination_7_15_2025.xlsx")
    st.success(f"✅ Excel file loaded successfully! Found {len(df_original)} records.")
    
except Exception as e:
    st.error(f"❌ Error loading Excel file: {e}")
    st.info("💡 Make sure 'sbd first_submission_clean.xlsx' is in the same directory as this app")
    st.stop()

# Load shapefile (embedded)
try:
    gdf = gpd.read_file("Chiefdom2021.shp")
    st.success(f"✅ Shapefile loaded successfully! Found {len(gdf)} features.")
    
except Exception as e:
    st.error(f"❌ Could not load shapefile: {e}")
    st.info("💡 Make sure 'Chiefdom2021.shp' and supporting files (.dbf, .shx, .prj) are in the same directory as this app")
    st.stop()

# Dashboard Settings - Fixed configuration
columns = 4  # Fixed to 4 columns for optimal Word export
show_itn_details = True  # Always show ITN data details

# Extract ITN coverage data
try:
    itn_df = extract_itn_data_from_excel(df_original)
    st.success(f"✅ ITN data extracted successfully! Found {len(itn_df)} records.")
except Exception as e:
    st.error(f"Error extracting ITN data: {e}")
    itn_df = pd.DataFrame()

if show_itn_details and len(itn_df) > 0:
    # Display ITN data information
    st.subheader("📊 ITN Data Overview")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_enrollment = int(itn_df["Total_Enrollment"].sum())
        st.metric("Total Enrollment", f"{total_enrollment:,}")
    
    with col2:
        total_distributed = int(itn_df["Distributed_ITNs"].sum())
        st.metric("ITNs Distributed", f"{total_distributed:,}")
    
    with col3:
        overall_coverage = (total_distributed / total_enrollment * 100) if total_enrollment > 0 else 0
        st.metric("Overall Coverage", f"{overall_coverage:.1f}%")
    
    with col4:
        schools_with_distribution = len(itn_df[itn_df["Distributed_ITNs"] > 0])
        st.metric("Schools with ITNs", f"{schools_with_distribution:,}")

# Create dashboards
st.header("🛡️ ITN Coverage Dashboards")

# BO District ITN Coverage Dashboard
st.subheader("BO District - ITN Coverage")

with st.spinner("Generating BO District ITN coverage dashboard..."):
    try:
        fig_bo_itn = create_itn_coverage_dashboard(gdf, itn_df, "BO", columns)
        if fig_bo_itn:
            st.pyplot(fig_bo_itn)
            
            # Save figure option
            buffer_bo_itn = BytesIO()
            fig_bo_itn.savefig(buffer_bo_itn, format='png', dpi=300, bbox_inches='tight')
            buffer_bo_itn.seek(0)
            
            st.download_button(
                label="📥 Download BO District ITN Coverage Dashboard (PNG)",
                data=buffer_bo_itn,
                file_name="BO_District_ITN_Coverage_Dashboard.png",
                mime="image/png"
            )
            
            # Word Export for BO District
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Create Word document
                doc = Document()
                
                # Add title
                title = doc.add_heading('BO District - ITN Coverage Analysis', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add generation date
                date_para = doc.add_paragraph()
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                date_run = date_para.add_run(f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
                date_run.font.size = Pt(12)
                
                doc.add_paragraph()  # Add space
                
                # Save matplotlib figure as PNG and embed in Word
                timestamp = pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')
                png_filename = f"BO_District_ITN_Coverage_{timestamp}.png"
                
                # Save PNG file to current directory
                fig_bo_itn.savefig(png_filename, format='png', dpi=200, 
                                  bbox_inches='tight', facecolor='white', 
                                  edgecolor='none', pad_inches=0.1)
                
                # Add the saved PNG to Word document
                doc.add_picture(png_filename, width=Inches(9.5))  # Fits well in Word page
                
                # Add format explanation
                doc.add_heading('ITN Coverage Format', level=2)
                format_items = [
                    "Title: (ITNs Distributed, Total Enrollment)",
                    "Center: Coverage percentage",
                    "ITN Calculation: Boys ITNs + Girls ITNs + ITNs Left at School",
                    "🔴 Red: < 20% coverage",
                    "🟠 Orange: 20-39% coverage", 
                    "🟡 Yellow: 40-59% coverage",
                    "🟢 Light Green: 60-79% coverage",
                    "🔵 Blue: 80-99% coverage",
                    "🟣 Purple: 100%+ coverage"
                ]
                
                for item in format_items:
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(item)
                
                # Add summary information
                doc.add_heading('Dashboard Summary', level=2)
                
                bo_data = itn_df[itn_df["District"].str.upper() == "BO"]
                bo_enrollment = int(bo_data["Total_Enrollment"].sum())
                bo_distributed = int(bo_data["Distributed_ITNs"].sum())
                bo_coverage = (bo_distributed / bo_enrollment * 100) if bo_enrollment > 0 else 0
                
                summary_text = f"""
                District: BO
                Total Chiefdoms: {len(gdf[gdf['FIRST_DNAM'] == 'BO'])}
                Total Enrollment: {bo_enrollment:,}
                ITNs Distributed: {bo_distributed:,}
                ITN Coverage: {bo_coverage:.1f}%
                PNG File Saved: {png_filename}
                """
                
                for line in summary_text.strip().split('\n'):
                    if line.strip():
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(line.strip())
                
                # Save to BytesIO
                word_buffer = BytesIO()
                doc.save(word_buffer)
                word_data = word_buffer.getvalue()
                
                # Success message
                st.success(f"✅ PNG saved as: {png_filename}")
                
                st.download_button(
                    label="📄 Download BO District ITN Report (Word)",
                    data=word_data,
                    file_name=f"BO_District_ITN_Report_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except ImportError:
                st.warning("⚠️ Word export requires python-docx library. Install with: pip install python-docx")
            except Exception as e:
                st.warning(f"⚠️ Word export failed: {str(e)}")
        else:
            st.warning("Could not generate BO District ITN coverage dashboard")
    except Exception as e:
        st.error(f"Error generating BO District ITN coverage dashboard: {e}")

st.divider()

# BOMBALI District ITN Coverage Dashboard
st.subheader("BOMBALI District - ITN Coverage")

with st.spinner("Generating BOMBALI District ITN coverage dashboard..."):
    try:
        fig_bombali_itn = create_itn_coverage_dashboard(gdf, itn_df, "BOMBALI", columns)
        if fig_bombali_itn:
            st.pyplot(fig_bombali_itn)
            
            # Save figure option
            buffer_bombali_itn = BytesIO()
            fig_bombali_itn.savefig(buffer_bombali_itn, format='png', dpi=300, bbox_inches='tight')
            buffer_bombali_itn.seek(0)
            
            st.download_button(
                label="📥 Download BOMBALI District ITN Coverage Dashboard (PNG)",
                data=buffer_bombali_itn,
                file_name="BOMBALI_District_ITN_Coverage_Dashboard.png",
                mime="image/png"
            )
            
            # Word Export for BOMBALI District
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Create Word document
                doc = Document()
                
                # Add title
                title = doc.add_heading('BOMBALI District - ITN Coverage Analysis', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add generation date
                date_para = doc.add_paragraph()
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                date_run = date_para.add_run(f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
                date_run.font.size = Pt(12)
                
                doc.add_paragraph()  # Add space
                
                # Save matplotlib figure as PNG and embed in Word
                timestamp = pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')
                png_filename = f"BOMBALI_District_ITN_Coverage_{timestamp}.png"
                
                # Save PNG file to current directory
                fig_bombali_itn.savefig(png_filename, format='png', dpi=200, 
                                       bbox_inches='tight', facecolor='white', 
                                       edgecolor='none', pad_inches=0.1)
                
                # Add the saved PNG to Word document
                doc.add_picture(png_filename, width=Inches(9.5))  # Fits well in Word page
                
                # Add format explanation
                doc.add_heading('ITN Coverage Format', level=2)
                format_items = [
                    "Title: (ITNs Distributed, Total Enrollment)",
                    "Center: Coverage percentage",
                    "ITN Calculation: Boys ITNs + Girls ITNs + ITNs Left at School",
                    "🔴 Red: < 20% coverage",
                    "🟠 Orange: 20-39% coverage", 
                    "🟡 Yellow: 40-59% coverage",
                    "🟢 Light Green: 60-79% coverage",
                    "🔵 Blue: 80-99% coverage",
                    "🟣 Purple: 100%+ coverage"
                ]
                
                for item in format_items:
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(item)
                
                # Add summary information
                doc.add_heading('Dashboard Summary', level=2)
                
                bombali_data = itn_df[itn_df["District"].str.upper() == "BOMBALI"]
                bombali_enrollment = int(bombali_data["Total_Enrollment"].sum())
                bombali_distributed = int(bombali_data["Distributed_ITNs"].sum())
                bombali_coverage = (bombali_distributed / bombali_enrollment * 100) if bombali_enrollment > 0 else 0
                
                summary_text = f"""
                District: BOMBALI
                Total Chiefdoms: {len(gdf[gdf['FIRST_DNAM'] == 'BOMBALI'])}
                Total Enrollment: {bombali_enrollment:,}
                ITNs Distributed: {bombali_distributed:,}
                ITN Coverage: {bombali_coverage:.1f}%
                PNG File Saved: {png_filename}
                """
                
                for line in summary_text.strip().split('\n'):
                    if line.strip():
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(line.strip())
                
                # Save to BytesIO
                word_buffer = BytesIO()
                doc.save(word_buffer)
                word_data = word_buffer.getvalue()
                
                # Success message
                st.success(f"✅ PNG saved as: {png_filename}")
                
                st.download_button(
                    label="📄 Download BOMBALI District ITN Report (Word)",
                    data=word_data,
                    file_name=f"BOMBALI_District_ITN_Report_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except ImportError:
                st.warning("⚠️ Word export requires python-docx library. Install with: pip install python-docx")
            except Exception as e:
                st.warning(f"⚠️ Word export failed: {str(e)}")
        else:
            st.warning("Could not generate BOMBALI District ITN coverage dashboard")
    except Exception as e:
        st.error(f"Error generating BOMBALI District ITN coverage dashboard: {e}")

# Export All Dashboards as Combined Word Document
st.header("📄 Combined Word Export")

if st.button("📋 Generate Combined ITN Report", help="Generate a comprehensive Word document with both districts"):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create Word document
        doc = Document()
        
        # Add main title
        title = doc.add_heading('School-Based Distribution (SBD)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('ITN Coverage Analysis Dashboard', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_run.font.size = Pt(12)
        date_run.bold = True
        
        doc.add_page_break()
        
        # ITN Coverage Format Explanation
        doc.add_heading('ITN Coverage Format', level=1)
        format_items = [
            "Title Format: (ITNs Distributed, Total Enrollment)",
            "Center Display: Coverage percentage",
            "ITN Calculation: Boys ITNs + Girls ITNs + ITNs Left at School",
            "🔴 Red: < 20% coverage (Critical - urgent action needed)",
            "🟠 Orange: 20-39% coverage (Poor - requires immediate attention)", 
            "🟡 Yellow: 40-59% coverage (Fair - needs improvement)",
            "🟢 Light Green: 60-79% coverage (Good - meeting targets)",
            "🔵 Blue: 80-99% coverage (Excellent - exceeding expectations)",
            "🟣 Purple: 100%+ coverage (Outstanding - full coverage achieved)"
        ]
        
        for item in format_items:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(item)
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        
        total_enrollment = int(itn_df["Total_Enrollment"].sum())
        total_distributed = int(itn_df["Distributed_ITNs"].sum())
        overall_coverage = (total_distributed / total_enrollment * 100) if total_enrollment > 0 else 0
        
        bo_data = itn_df[itn_df["District"].str.upper() == "BO"]
        bombali_data = itn_df[itn_df["District"].str.upper() == "BOMBALI"]
        
        bo_enrollment = int(bo_data["Total_Enrollment"].sum())
        bo_distributed = int(bo_data["Distributed_ITNs"].sum())
        bo_coverage = (bo_distributed / bo_enrollment * 100) if bo_enrollment > 0 else 0
        
        bombali_enrollment = int(bombali_data["Total_Enrollment"].sum())
        bombali_distributed = int(bombali_data["Distributed_ITNs"].sum())
        bombali_coverage = (bombali_distributed / bombali_enrollment * 100) if bombali_enrollment > 0 else 0
        
        summary_text = f"""
        This comprehensive dashboard report presents ITN distribution effectiveness analysis for BO and BOMBALI districts:
        
        • Districts Covered: BO, BOMBALI
        • Total Student Enrollment: {total_enrollment:,}
        • Total ITNs Distributed: {total_distributed:,}
        • Overall ITN Coverage: {overall_coverage:.1f}%
        • BO District Coverage: {bo_coverage:.1f}%
        • BOMBALI District Coverage: {bombali_coverage:.1f}%
        
        Coverage is calculated as: (ITNs Distributed / Total Enrollment) × 100%
        
        Formulas:
        • ITN Distributed = Boys ITNs + Girls ITNs + ITNs Left at School
        • Total Enrollment = Sum of all pupils enrolled in Classes 1-5
        • Coverage = (ITNs Distributed / Total Enrollment) × 100%
        
        Color coding helps identify areas requiring attention and those performing well.
        Format shows (ITNs Distributed, Total Enrollment) with coverage percentage in center.
        """
        
        for line in summary_text.strip().split('\n'):
            if line.strip():
                if line.startswith('•'):
                    p = doc.add_paragraph()
                    p.add_run(line.strip())
                else:
                    doc.add_paragraph(line.strip())
        
        doc.add_page_break()
        
        # BO District section
        if 'fig_bo_itn' in locals():
            doc.add_heading('BO District - ITN Coverage Analysis', level=1)
            
            # Save BO figure as PNG and embed in Word
            timestamp = pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')
            bo_png_filename = f"BO_District_ITN_Combined_{timestamp}.png"
            
            # Save PNG file to current directory
            fig_bo_itn.savefig(bo_png_filename, format='png', dpi=200, 
                              bbox_inches='tight', facecolor='white', 
                              edgecolor='none', pad_inches=0.1)
            
            # Add the saved PNG to Word document
            doc.add_picture(bo_png_filename, width=Inches(9.5))  # Fits well in Word page
            
            # BO summary
            doc.add_heading('BO District Summary', level=2)
            
            bo_summary_items = [
                f"Total Chiefdoms: {len(gdf[gdf['FIRST_DNAM'] == 'BO'])}",
                f"Total Enrollment: {bo_enrollment:,}",
                f"ITNs Distributed: {bo_distributed:,}",
                f"ITN Coverage: {bo_coverage:.1f}%",
                f"PNG File Saved: {bo_png_filename}"
            ]
            
            for item in bo_summary_items:
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(item)
            
            doc.add_page_break()
        
        # BOMBALI District section
        if 'fig_bombali_itn' in locals():
            doc.add_heading('BOMBALI District - ITN Coverage Analysis', level=1)
            
            # Save BOMBALI figure as PNG and embed in Word
            bombali_png_filename = f"BOMBALI_District_ITN_Combined_{timestamp}.png"
            
            # Save PNG file to current directory
            fig_bombali_itn.savefig(bombali_png_filename, format='png', dpi=200, 
                                   bbox_inches='tight', facecolor='white', 
                                   edgecolor='none', pad_inches=0.1)
            
            # Add the saved PNG to Word document
            doc.add_picture(bombali_png_filename, width=Inches(9.5))  # Fits well in Word page
            
            # BOMBALI summary
            doc.add_heading('BOMBALI District Summary', level=2)
            
            bombali_summary_items = [
                f"Total Chiefdoms: {len(gdf[gdf['FIRST_DNAM'] == 'BOMBALI'])}",
                f"Total Enrollment: {bombali_enrollment:,}",
                f"ITNs Distributed: {bombali_distributed:,}",
                f"ITN Coverage: {bombali_coverage:.1f}%",
                f"PNG File Saved: {bombali_png_filename}"
            ]
            
            for item in bombali_summary_items:
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(item)
        
        # Save to BytesIO
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_data = word_buffer.getvalue()
        
        # Success message showing saved PNG files
        st.success(f"✅ PNG files saved: {bo_png_filename}, {bombali_png_filename}")
        
        st.download_button(
            label="💾 Download Combined ITN Analysis Report (Word)",
            data=word_data,
            file_name=f"ITN_Coverage_Analysis_Report_{timestamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Download comprehensive Word report with both districts"
        )
        
    except ImportError:
        st.error("❌ Word generation requires python-docx library. Please install it using: pip install python-docx")
    except Exception as e:
        st.error(f"❌ Error generating combined Word document: {str(e)}")

# ITN Analysis
st.header("📈 ITN Distribution Analysis")

if len(itn_df) > 0:
    # Calculate ITN statistics
    bo_data = itn_df[itn_df["District"].str.upper() == "BO"]
    bombali_data = itn_df[itn_df["District"].str.upper() == "BOMBALI"]
    
    # ITN metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        bo_enrollment = int(bo_data["Total_Enrollment"].sum())
        bo_distributed = int(bo_data["Distributed_ITNs"].sum())
        bo_coverage = (bo_distributed / bo_enrollment * 100) if bo_enrollment > 0 else 0
        st.metric("BO District ITN Coverage", f"{bo_coverage:.1f}%", f"{bo_distributed}/{bo_enrollment}")
    
    with col2:
        bombali_enrollment = int(bombali_data["Total_Enrollment"].sum())
        bombali_distributed = int(bombali_data["Distributed_ITNs"].sum())
        bombali_coverage = (bombali_distributed / bombali_enrollment * 100) if bombali_enrollment > 0 else 0
        st.metric("BOMBALI District ITN Coverage", f"{bombali_coverage:.1f}%", f"{bombali_distributed}/{bombali_enrollment}")
    
    with col3:
        total_enrollment = int(itn_df["Total_Enrollment"].sum())
        total_distributed = int(itn_df["Distributed_ITNs"].sum())
        overall_coverage = (total_distributed / total_enrollment * 100) if total_enrollment > 0 else 0
        st.metric("Overall ITN Coverage", f"{overall_coverage:.1f}%", f"{total_distributed}/{total_enrollment}")
    
    with col4:
        # Calculate chiefdoms with good ITN coverage (>= 60%)
        good_itn_coverage_count = 0
        total_chiefdoms = 0
        
        for district in ["BO", "BOMBALI"]:
            district_data = itn_df[itn_df["District"].str.upper() == district.upper()]
            chiefdoms = district_data['Chiefdom'].dropna().unique()
            
            for chiefdom in chiefdoms:
                chiefdom_data = district_data[district_data['Chiefdom'] == chiefdom]
                enrollment = int(chiefdom_data["Total_Enrollment"].sum())
                distributed = int(chiefdom_data["Distributed_ITNs"].sum())
                coverage = (distributed / enrollment * 100) if enrollment > 0 else 0
                
                if coverage >= 60:
                    good_itn_coverage_count += 1
                total_chiefdoms += 1
        
        good_itn_coverage_percent = (good_itn_coverage_count / total_chiefdoms * 100) if total_chiefdoms > 0 else 0
        st.metric("Chiefdoms with Good ITN Coverage", f"{good_itn_coverage_percent:.0f}%", f"{good_itn_coverage_count}/{total_chiefdoms}")
    
    # Detailed ITN coverage table
    st.subheader("📋 Detailed ITN Coverage by Chiefdom")
    
    detailed_itn_coverage = []
    for district in ["BO", "BOMBALI"]:
        district_data = itn_df[itn_df["District"].str.upper() == district.upper()]
        chiefdoms = sorted(district_data['Chiefdom'].dropna().unique())
        
        for chiefdom in chiefdoms:
            chiefdom_data = district_data[district_data['Chiefdom'] == chiefdom]
            enrollment = int(chiefdom_data["Total_Enrollment"].sum())
            distributed = int(chiefdom_data["Distributed_ITNs"].sum())
            coverage = (distributed / enrollment * 100) if enrollment > 0 else 0
            
            # Determine status
            if coverage >= 80:
                status = "✅ Excellent"
            elif coverage >= 60:
                status = "🟢 Good"
            elif coverage >= 40:
                status = "🟡 Fair"
            elif coverage >= 20:
                status = "🟠 Poor"
            else:
                status = "🔴 Critical"
            
            detailed_itn_coverage.append({
                'District': district,
                'Chiefdom': chiefdom,
                'Total Enrollment': enrollment,
                'ITNs Distributed': distributed,
                'Coverage %': f"{coverage:.1f}%",
                'Status': status
            })
    
    itn_coverage_df = pd.DataFrame(detailed_itn_coverage)
    st.dataframe(itn_coverage_df, use_container_width=True)
    
    # Distribution Summary
    st.subheader("📊 Distribution Summary")
    
    try:
        summary_data = generate_simple_summary(itn_df)
        summary_df = pd.DataFrame(summary_data)
        
        st.dataframe(summary_df, use_container_width=True)
    
    except Exception as e:
        st.error(f"Error generating distribution summary: {e}")

# Memory optimization - close matplotlib figures
plt.close('all')

# Footer
st.markdown("---")
st.markdown("**🛡️ Section 3: ITN Coverage Analysis | School-Based Distribution Analysis**")
